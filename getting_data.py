import pdfplumber, re, docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL

SUBJECT_MAP = {
    'TOÁN': 'Toán', 'VĂN': 'Ngữ Văn', 'ANH': 'Tiếng Anh',
    'LÝ': 'Vật Lý', 'HÓA': 'Hóa Học', 'SINH': 'Sinh Học',
    'SỬ': 'Lịch Sử', 'ĐỊA': 'Địa Lý', 'TIN': 'Tin Học',
    'GDTC': 'GDTC', 'GDQP': 'GDQP', 'CN': 'Công Nghệ',
    'CNCN': 'Công Nghệ', 'CNNN': 'Công Nghệ', 
    'KTPL': 'GDKT&PL', 'HĐTN': 'HĐTN', 'GDĐP': 'GDĐP',
    'SHL': 'Sinh Hoạt', 'HĐSHDC': 'Chào cờ', 'GDKT-PL': 'GDKT&PL'
}

def clean_subject_name(raw_text):
    if not raw_text or raw_text.strip() == "": return "—"
    first_line = raw_text.split('\n')[0].strip()
    if "10G" in first_line.upper(): return None
    prefix = first_line.split('-')[0].strip().upper()
    return SUBJECT_MAP.get(prefix, first_line)

def get_data_from_pdf(pdf_path):
    schedule_data = {day: [] for day in range(2, 8)}
    apply_date = "dd/mm/yyyy"
    with pdfplumber.open(pdf_path) as pdf:
        page = pdf.pages[0]
        text = page.extract_text()
        date_match = re.search(r'NGÀY (\d{2}[-/]\d{2}[-/]\d{4})', text, re.IGNORECASE)
        if date_match:
            apply_date = date_match.group(1).replace('-', '/')
        
        table = page.extract_tables()[0]
        class_col_idx = 8
        current_day = 2
        for row in table[1:]:
            if not any(row): continue
            thu_val = str(row[0]).strip()
            if thu_val.isdigit() and 2 <= int(thu_val) <= 7:
                current_day = int(thu_val)
            raw_subj = str(row[class_col_idx]).strip() if row[class_col_idx] else ""
            cleaned = clean_subject_name(raw_subj)
            if cleaned is not None:
                if len(schedule_data[current_day]) < 5:
                    schedule_data[current_day].append(cleaned)
    return schedule_data, apply_date

def set_run_format(run, is_dash=False):
    run.font.name = 'Mulish'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Mulish')
    if is_dash:
        run.font.size = Pt(12)
        run.bold = True
    else:
        pass 

def update_cell_precise(cell, new_text):
    if not cell.paragraphs:
        cell.add_paragraph()
    para = cell.paragraphs[0]
    is_dash = (new_text.strip() == "—")
    if para.runs:
        target_run = para.runs[0]
        target_run.text = new_text
        set_run_format(target_run, is_dash)
        for i in range(1, len(para.runs)):
            para.runs[i].text = ""
    else:
        new_run = para.add_run(new_text)
        set_run_format(new_run, is_dash)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def merge_identical_cells(table, row_indices, col_range):
    for col in col_range:
        r = 0
        while r < len(row_indices) - 1:
            curr_row = row_indices[r]
            curr_cell = table.cell(curr_row, col)
            curr_text = curr_cell.text.strip()
            
            if curr_text == "—" or not curr_text:
                r += 1
                continue
            
            # Tìm danh sách các ô tiếp theo có cùng nội dung trong cùng một cột
            match_rows = []
            for next_idx in range(r + 1, len(row_indices)):
                next_row = row_indices[next_idx]
                if table.cell(next_row, col).text.strip() == curr_text:
                    match_rows.append(next_row)
                else:
                    break
            
            if match_rows:
                # 1. Xóa nội dung các ô trùng trước khi gộp
                for row_to_clear in match_rows:
                    cell_to_clear = table.cell(row_to_clear, col)
                    for p in cell_to_clear.paragraphs:
                        for run in p.runs:
                            run.text = ""
                        p.text = ""

                # 2. Thực hiện gộp ô đầu tiên với ô trùng cuối cùng
                last_row = match_rows[-1]
                merged_cell = table.cell(curr_row, col).merge(table.cell(last_row, col))
                
                # 3. Xóa bỏ các paragraph thừa phát sinh do lệnh merge
                while len(merged_cell.paragraphs) > 1:
                    p_element = merged_cell.paragraphs[-1]._element
                    p_element.getparent().remove(p_element)
                
                # 4. Ghi lại nội dung chuẩn vào ô đã gộp duy nhất
                update_cell_precise(merged_cell, curr_text)
                
                r += len(match_rows) + 1
            else:
                r += 1

def fill_docx(data, date_str, TEMPLATE_PATH, OUTPUT_PATH):
    doc = docx.Document(TEMPLATE_PATH)
    
    # 1. Sửa Bảng 1 (Tiêu đề)
    table_header = doc.tables[0]
    for row in table_header.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if "dd/mm/yyyy" in para.text:
                    full_text = para.text.replace("dd/mm/yyyy", date_str)
                    if para.runs:
                        para.runs[0].text = full_text
                        for i in range(1, len(para.runs)):
                            para.runs[i].text = ""
                    else:
                        para.add_run(full_text)

    # 2. Sửa Bảng 2: Điền TKB Sáng
    tkb_table = doc.tables[1]
    row_indices = [3, 4, 5, 6, 7]
    col_range = range(2, 8) 
    
    # Bước A: Điền dữ liệu thô vào từng ô
    for day in range(2, 8):
        col_idx = day
        subjects = data[day]
        for i, row_idx in enumerate(row_indices):
            val = subjects[i] if i < len(subjects) else "—"
            if row_idx < len(tkb_table.rows) and col_idx < len(tkb_table.columns):
                cell = tkb_table.cell(row_idx, col_idx)
                update_cell_precise(cell, val)

    # Bước B: Thực hiện logic gộp ô đã cải tiến
    merge_identical_cells(tkb_table, row_indices, col_range)

    doc.save(OUTPUT_PATH)

def start(PDF_PATH, TEMPLATE_PATH, OUTPUT_PATH):
    data, date_val = get_data_from_pdf(PDF_PATH)
    fill_docx(data, date_val, TEMPLATE_PATH, OUTPUT_PATH)
    print(f"Tạo TKB.docx thành công!\n")